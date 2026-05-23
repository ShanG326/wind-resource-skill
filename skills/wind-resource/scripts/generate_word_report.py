#!/usr/bin/env python3
"""Generate Word report from wind resource assessment data.

Modify CHARTS_DIR, OUTPUT_DIR, and data dictionaries to match your project.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


OUTPUT_DIR = os.path.join(os.path.expanduser("~"), "Desktop")
CHARTS_DIR = ""  # Set to your charts directory


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shd)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return h


def add_para(doc, text, bold=False, size=Pt(10.5)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = size
    run.font.bold = bold
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.name = 'SimSun'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        set_cell_shading(cell, '003366')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'SimSun'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        if r_idx % 2 == 1:
            for c_idx in range(len(headers)):
                set_cell_shading(row.cells[c_idx], 'E8EDF2')
    doc.add_paragraph()
    return table


def add_chart(doc, image_path, width_inches=5.5):
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(width_inches))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_para(doc, f"[Missing image: {os.path.basename(image_path)}]")


def add_tower_report(doc, tower_id, charts_dir, data):
    add_heading(doc, f"Mast {tower_id} Wind Resource Assessment", level=1)
    add_heading(doc, "1. Project Overview", level=2)
    add_heading(doc, "1.1 Standards", level=3)
    for item in data.get("standards", []):
        add_bullet(doc, item)
    add_heading(doc, "1.2 Basic Info", level=3)
    add_para(doc, f"Mast ID: {tower_id}")
    add_para(doc, f"Period: {data.get('data_start', '')} ~ {data.get('data_end', '')}")
    add_para(doc, "Interval: 10 min")

    add_heading(doc, "2. Data Overview", level=2)
    add_para(doc, f"Total records: {data.get('total_records', '')}")
    add_para(doc, f"Valid records: {data.get('valid_records', '')}")
    add_para(doc, f"Completeness: {data.get('completeness', '')}")
    add_para(doc, f"Coverage months: {data.get('coverage_months', '')}")

    add_heading(doc, "3. Wind Speed Analysis", level=2)
    add_para(doc, f"Mean temperature: {data.get('avg_temp', '')} C")
    add_para(doc, f"Mean pressure: {data.get('avg_pressure', '')} kPa")
    add_para(doc, f"Mean air density: {data.get('air_density', '')} kg/m3")

    if charts_dir and os.path.isdir(charts_dir):
        for fn in sorted(os.listdir(charts_dir)):
            if fn.endswith('.png'):
                add_chart(doc, os.path.join(charts_dir, fn), 4.5)

    add_heading(doc, "4. Conclusions", level=2)
    for item in data.get("conclusions", []):
        add_para(doc, item)

    doc.add_page_break()


def main():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    # Cover page
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Wind Resource Assessment Report")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    doc.add_page_break()

    # Add tower reports here
    # Example:
    # add_tower_report(doc, "7152", CHARTS_DIR, data_7152)

    output_path = os.path.join(OUTPUT_DIR, "wind_resource_report.docx")
    doc.save(output_path)
    print(f"Word report saved: {output_path}")


if __name__ == "__main__":
    main()
